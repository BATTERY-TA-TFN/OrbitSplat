from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "course_report"
OUT_DIR.mkdir(parents=True, exist_ok=True)

REPORT_DOCX = OUT_DIR / "基于Gaussian_Splatting的真实物体三维重建课程报告.docx"
REPORT_MD = OUT_DIR / "基于Gaussian_Splatting的真实物体三维重建课程报告.md"

FIGURES = {
    "keyframes": ROOT / "output" / "keyframe_selection_figure.png",
    "filtering": ROOT / "output" / "colmap_mask_filtering_figure.png",
    "init_chart": ROOT / "output" / "evaluation" / "initialization_ablation_metrics.png",
    "init_qual": ROOT / "output" / "evaluation" / "initialization_ablation_qualitative.png",
    "view_chart": ROOT / "output" / "evaluation" / "common_subset_metrics.png",
    "view_qual": ROOT / "output" / "evaluation" / "common_subset_qualitative.png",
    "repair_chart": ROOT / "output" / "evaluation" / "repair_engineering_validation.png",
}


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, True)
        set_cell_shading(table.rows[0].cells[index], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
    document.add_paragraph()
    return table


def add_picture(document, path, caption, width_cm=15.2):
    if path.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(path), width=Cm(width_cm))
        cap = document.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(90, 90, 90)


def add_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    return paragraph


def add_paragraph(document, text):
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.line_spacing = 1.25
    for run in paragraph.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)
    return paragraph


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)
        for run in paragraph.runs:
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(10.5)


def add_code(document, code):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.right_indent = Pt(18)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def build_docx():
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("基于 Gaussian Splatting 的真实物体三维重建设计与实现")
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("课程报告")
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.name = "黑体"
    subtitle.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    info = document.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("姓名：王婧    环境：Windows + RTX 3080    日期：2026 年 6 月")
    for run in info.runs:
        run.font.size = Pt(10.5)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    add_heading(document, "一、选题背景与意义", 1)
    add_paragraph(document, "随着数字孪生、AR/VR、虚拟展示和三维内容生成的发展，低成本、自动化的三维重建技术逐渐成为数字图像处理和计算机视觉领域的重要方向。传统三维建模依赖专业设备、人工建模或复杂扫描流程，成本较高；而 NeRF 与 3D Gaussian Splatting 等新视角合成技术使得仅利用二维图像重建三维物体成为可能。")
    add_paragraph(document, "本项目选择真实手机视频作为输入对象，目标不是只在标准数据集上复现论文结果，而是探索普通用户能否在 Windows 环境下使用普通手机拍摄的视频完成真实物体三维重建。该选题具有较强的应用意义：一方面降低三维内容制作门槛，另一方面也能暴露论文方法在真实采集环境中的数据质量、位姿估计、分割误差和工程部署问题。")

    add_heading(document, "二、项目环境与软件版本", 1)
    add_table(document, ["类别", "配置"], [
        ["操作系统", "Windows"],
        ["GPU", "NVIDIA RTX 3080"],
        ["Python", "3.11"],
        ["CUDA", "11.8"],
        ["PyTorch", "2.1.2"],
        ["主要工具", "COLMAP、OpenCV、GaussianObject、3D Gaussian Splatting、LPIPS"],
        ["数据来源", "手机环绕拍摄皮卡丘视频，共 54 个已标定位姿视角"],
    ])

    add_heading(document, "三、参考方法与总体技术路线", 1)
    add_paragraph(document, "本项目以 GaussianObject 为参考框架。GaussianObject 的核心思想是在少视图条件下利用 Visual Hull、粗糙 3DGS 表示和 Gaussian Repair 修复模型提高物体重建质量。我的工作重点并不是重新设计整套生成式修复网络，而是围绕真实手机视频输入，对数据预处理、物体分割、点云初始化、位姿估计和实验评价进行工程化适配与方法改进。")
    add_paragraph(document, "项目整体流程包括：手机环绕视频采集、关键帧提取与清晰度筛选、多阶段前景分割、COLMAP/DUSt3R/MASt3R 位姿路线比较、掩膜约束点云过滤、粗糙 3DGS 训练、Gaussian Repair 工程验证以及新视角渲染和视频生成。与直接运行论文代码相比，本项目需要处理真实视频中的模糊帧、重复帧、背景干扰、低纹理匹配失败和掩膜边缘误差等问题。")

    add_heading(document, "四、主要工作与个人贡献", 1)
    add_heading(document, "4.1 手机视频关键帧选择", 2)
    add_paragraph(document, "真实手机视频中存在大量相似帧和运动模糊帧。如果直接使用全部帧训练，会增加计算负担，也可能因低质量图像引入错误监督。因此我实现了关键帧筛选流程：首先使用拉普拉斯方差评价图像清晰度，再结合时间窗口和视角覆盖原则选择代表性帧，尽量保留侧面、背面和顶部等互补视角。")
    add_picture(document, FIGURES["keyframes"], "图 1 关键帧选择与清晰度筛选示意", 14.8)

    add_heading(document, "4.2 多阶段物体分割", 2)
    add_paragraph(document, "掩膜质量直接影响 Visual Hull、轮廓约束和点云过滤。为减少单一分割方法带来的不稳定性，我采用颜色先验、GrabCut/SAM 分割和形态学后处理相结合的方式生成前景掩膜。该过程先快速定位目标区域，再进行精细分割和连通域筛选，最后通过孔洞填充、边缘平滑等操作获得较稳定的二值掩膜。")

    add_heading(document, "4.3 多视图掩膜约束点云过滤", 2)
    add_paragraph(document, "这是本项目最重要的方法改进。真实桌面拍摄场景中，COLMAP 会重建出大量桌面和背景稀疏点，这些点会污染 3DGS 的初始几何。我将 COLMAP 稀疏点投影到多个相机视图中，统计每个点落入前景掩膜的支持次数，仅保留满足多视图前景一致性的点。")
    add_code(document, "for each 3D point X:\n    support = 0\n    visible = 0\n    for each camera view i:\n        u = project(X, camera_i)\n        if u is inside image:\n            visible += 1\n            if mask_i[u] == foreground:\n                support += 1\n    keep X if support / visible >= threshold")
    add_paragraph(document, "该方法的优点是利用多视图一致性而不是单张图像判断前景，因此能够减少偶然的掩膜误差；同时它直接作用于 COLMAP 稀疏点云，能够在训练 3DGS 之前降低背景污染。")
    add_picture(document, FIGURES["filtering"], "图 2 原始 COLMAP 点云、掩膜过滤点云与重建结果对比", 15.2)

    add_heading(document, "五、实验设计与结果分析", 1)
    add_paragraph(document, "实验数据为手机环绕拍摄的皮卡丘视频，共 54 个视角。严格少视角设置中选取 8 个训练视角，其余 46 个视角作为测试集；训练分辨率统一缩放为原图的 1/4，训练 10,000 次迭代。评价指标包括 PSNR、SSIM、LPIPS，并补充前景区域 PSNR 与前景裁剪 LPIPS。")

    add_heading(document, "5.1 初始化方式消融实验", 2)
    add_table(document, ["初始化方法", "初始点数", "PSNR↑", "SSIM↑", "LPIPS↓", "前景PSNR↑", "前景裁剪LPIPS↓"], [
        ["原始 COLMAP", "5,727", "26.596", "0.96273", "0.04459", "16.314", "0.26980"],
        ["掩膜过滤 COLMAP", "285", "26.519", "0.96276", "0.04515", "16.246", "0.26907"],
        ["Visual Hull", "7,787", "24.713", "0.95745", "0.05733", "13.288", "0.31887"],
    ])
    add_paragraph(document, "实验表明，多视图掩膜过滤将初始点数由 5,727 降至 285，仅保留约 4.98% 的点，去除了约 95.02% 的背景点。与此同时，PSNR 仅下降 0.077 dB，SSIM 略有提高，前景裁剪 LPIPS 也略有改善。这说明该方法能够显著减少背景污染，并基本保持重建质量。Visual Hull 在真实拍摄数据上表现较差，说明其对掩膜误差和位姿误差更加敏感。")
    add_picture(document, FIGURES["init_chart"], "图 3 初始化方式消融实验指标对比", 15.2)
    add_picture(document, FIGURES["init_qual"], "图 4 不同初始化方式在未见视角上的定性对比", 15.2)

    add_heading(document, "5.2 视角数量与覆盖范围实验", 2)
    add_table(document, ["方法", "PSNR↑", "SSIM↑", "LPIPS↓", "前景PSNR↑", "前景裁剪LPIPS↓"], [
        ["8 视角，包含少量顶部视角", "17.958", "0.92673", "0.10038", "7.522", "0.49354"],
        ["40 视角，主要为水平环绕视角", "17.414", "0.93085", "0.09961", "6.620", "0.50889"],
    ])
    add_paragraph(document, "为避免测试集不同造成误导，我选取 8 视角模型和 40 视角模型均未见过的 12 个顶部视角进行评价。结果表明，40 视角模型在 SSIM 和 LPIPS 上略好，但 8 视角模型在 PSNR 和前景指标上更高。该现象说明训练图像数量并不是唯一决定因素，视角的空间覆盖范围同样关键。如果训练视角主要集中在水平环绕，而缺少顶部视角，那么模型在顶部测试视角上仍会出现泛化不足。")
    add_picture(document, FIGURES["view_chart"], "图 5 8 视角与 40 视角在共同未见顶部视角上的指标对比", 15.2)
    add_picture(document, FIGURES["view_qual"], "图 6 共同未见顶部视角上的定性渲染结果", 15.2)

    add_heading(document, "5.3 Gaussian Repair 工程验证", 2)
    add_table(document, ["方法", "PSNR↑", "SSIM↑", "LPIPS↓"], [
        ["Coarse 3DGS", "28.295", "0.97198", "0.03663"],
        ["Gaussian Repair", "28.636", "0.97145", "0.03703"],
    ])
    add_paragraph(document, "该实验使用已有全视角 HQ 模型作为初始化，仅用于验证 Repair 链路是否能在本地真实数据上完整运行，不属于严格少视角留出集实验。结果显示 Gaussian Repair 使拟合视角 PSNR 提高 0.342 dB，但 SSIM 和 LPIPS 略有退化，定性变化也较小。因此，本项目没有将其表述为稳定提升的新结论，而是将其作为工程链路验证和后续改进方向。")
    add_picture(document, FIGURES["repair_chart"], "图 7 Gaussian Repair 工程验证指标对比", 15.2)

    add_heading(document, "六、项目难点与解决办法", 1)
    add_table(document, ["遇到的问题", "解决办法", "效果"], [
        ["视频帧模糊、重复度高", "使用拉普拉斯方差和时间窗口筛选关键帧", "减少低质量输入，降低训练负担"],
        ["背景点污染 COLMAP 点云", "提出多视图掩膜一致性点云过滤", "背景点减少约 95%，重建质量基本保持"],
        ["掩膜边缘误差影响几何", "使用颜色先验、SAM/GrabCut 与形态学后处理", "获得更稳定的前景掩膜"],
        ["真实物体纹理较少，位姿估计不稳定", "比较 COLMAP、DUSt3R、MASt3R 等路线", "确定更适合当前数据的位姿流程"],
        ["论文流程偏向 Linux 和标准数据集", "完成 Windows + RTX 3080 本地部署和脚本适配", "跑通真实手机视频到新视角渲染的完整链路"],
    ])

    add_heading(document, "七、收获与展望", 1)
    add_paragraph(document, "通过本项目，我不仅理解了 3D Gaussian Splatting 和 GaussianObject 的基本流程，也更深刻地认识到真实数据和标准数据集之间的差距。真实手机视频中的模糊、重复、反光、低纹理、背景干扰和掩膜误差都会被三维重建算法放大，因此工程实现并不是简单地运行论文代码，而需要围绕数据质量、位姿、掩膜和初始化做大量适配。")
    add_paragraph(document, "从实验结果看，本项目最有价值的结论是：多视图掩膜约束点云过滤能够显著去除背景污染，并在严格少视角测试中基本保持重建质量；同时，视角覆盖比单纯增加训练图像数量更加重要。后续工作可以进一步设计联合清晰度、相似度和空间覆盖的关键帧选择算法，并引入位姿置信度加权训练，从而提升真实场景下的重建稳定性。")

    add_heading(document, "八、结论", 1)
    add_paragraph(document, "本课程项目完成了基于 Gaussian Splatting 的真实物体三维重建设计与实现，构建了从手机视频采集、关键帧筛选、前景分割、位姿估计、点云过滤、3DGS 训练到新视角渲染的完整流程。项目在参考 GaussianObject 的基础上，重点解决真实手机视频输入下的工程适配问题，并提出多视图掩膜约束点云过滤方法。实验表明，该方法能够将 COLMAP 初始点云减少约 95%，同时基本保持渲染指标稳定，体现出较好的实用价值。")
    add_paragraph(document, "总体而言，本项目既完成了算法复现与工程部署，也通过多组消融实验分析了点云初始化、视角覆盖和修复策略对重建质量的影响，达到了课程报告对技术路线、实验结果和个人贡献展示的要求。")

    add_heading(document, "参考文献", 1)
    refs = [
        "Yang, Chen, et al. GaussianObject: High-Quality 3D Object Reconstruction from Four Views with Gaussian Splatting. ACM Transactions on Graphics, 2024.",
        "Kerbl, Bernhard, et al. 3D Gaussian Splatting for Real-Time Radiance Field Rendering. ACM Transactions on Graphics, 2023.",
        "Schonberger, Johannes L., and Jan-Michael Frahm. Structure-from-Motion Revisited. CVPR, 2016.",
    ]
    for ref in refs:
        document.add_paragraph(ref, style="List Number")

    document.save(REPORT_DOCX)


def build_markdown():
    text = f"""# 基于 Gaussian Splatting 的真实物体三维重建设计与实现

## 一、选题背景与意义

随着数字孪生、AR/VR、虚拟展示和三维内容生成的发展，低成本、自动化的三维重建技术逐渐成为数字图像处理和计算机视觉领域的重要方向。传统三维建模依赖专业设备、人工建模或复杂扫描流程，成本较高；而 NeRF 与 3D Gaussian Splatting 等新视角合成技术使得仅利用二维图像重建三维物体成为可能。

本项目选择真实手机视频作为输入对象，目标不是只在标准数据集上复现论文结果，而是探索普通用户能否在 Windows 环境下使用普通手机拍摄的视频完成真实物体三维重建。该选题具有较强的应用意义：一方面降低三维内容制作门槛，另一方面也能暴露论文方法在真实采集环境中的数据质量、位姿估计、分割误差和工程部署问题。

## 二、项目环境与软件版本

| 类别 | 配置 |
|---|---|
| 操作系统 | Windows |
| GPU | NVIDIA RTX 3080 |
| Python | 3.11 |
| CUDA | 11.8 |
| PyTorch | 2.1.2 |
| 主要工具 | COLMAP、OpenCV、GaussianObject、3D Gaussian Splatting、LPIPS |
| 数据来源 | 手机环绕拍摄皮卡丘视频，共 54 个已标定位姿视角 |

## 三、参考方法与总体技术路线

本项目以 GaussianObject 为参考框架。GaussianObject 的核心思想是在少视图条件下利用 Visual Hull、粗糙 3DGS 表示和 Gaussian Repair 修复模型提高物体重建质量。我的工作重点并不是重新设计整套生成式修复网络，而是围绕真实手机视频输入，对数据预处理、物体分割、点云初始化、位姿估计和实验评价进行工程化适配与方法改进。

项目整体流程包括：手机环绕视频采集、关键帧提取与清晰度筛选、多阶段前景分割、COLMAP/DUSt3R/MASt3R 位姿路线比较、掩膜约束点云过滤、粗糙 3DGS 训练、Gaussian Repair 工程验证以及新视角渲染和视频生成。

## 四、主要工作与个人贡献

### 4.1 手机视频关键帧选择

真实手机视频中存在大量相似帧和运动模糊帧。如果直接使用全部帧训练，会增加计算负担，也可能因低质量图像引入错误监督。因此我实现了关键帧筛选流程：首先使用拉普拉斯方差评价图像清晰度，再结合时间窗口和视角覆盖原则选择代表性帧。

![关键帧选择]({FIGURES["keyframes"]})

### 4.2 多阶段物体分割

掩膜质量直接影响 Visual Hull、轮廓约束和点云过滤。为减少单一分割方法带来的不稳定性，我采用颜色先验、GrabCut/SAM 分割和形态学后处理相结合的方式生成前景掩膜。

### 4.3 多视图掩膜约束点云过滤

这是本项目最重要的方法改进。真实桌面拍摄场景中，COLMAP 会重建出大量桌面和背景稀疏点，这些点会污染 3DGS 的初始几何。我将 COLMAP 稀疏点投影到多个相机视图中，统计每个点落入前景掩膜的支持次数，仅保留满足多视图前景一致性的点。

```python
for each 3D point X:
    support = 0
    visible = 0
    for each camera view i:
        u = project(X, camera_i)
        if u is inside image:
            visible += 1
            if mask_i[u] == foreground:
                support += 1
    keep X if support / visible >= threshold
```

![点云过滤]({FIGURES["filtering"]})

## 五、实验设计与结果分析

实验数据为手机环绕拍摄的皮卡丘视频，共 54 个视角。严格少视角设置中选取 8 个训练视角，其余 46 个视角作为测试集；训练分辨率统一缩放为原图的 1/4，训练 10,000 次迭代。评价指标包括 PSNR、SSIM、LPIPS，并补充前景区域 PSNR 与前景裁剪 LPIPS。

### 5.1 初始化方式消融实验

| 初始化方法 | 初始点数 | PSNR↑ | SSIM↑ | LPIPS↓ | 前景PSNR↑ | 前景裁剪LPIPS↓ |
|---|---:|---:|---:|---:|---:|---:|
| 原始 COLMAP | 5,727 | 26.596 | 0.96273 | 0.04459 | 16.314 | 0.26980 |
| 掩膜过滤 COLMAP | 285 | 26.519 | 0.96276 | 0.04515 | 16.246 | 0.26907 |
| Visual Hull | 7,787 | 24.713 | 0.95745 | 0.05733 | 13.288 | 0.31887 |

多视图掩膜过滤将初始点数由 5,727 降至 285，仅保留约 4.98% 的点，去除了约 95.02% 的背景点。与此同时，PSNR 仅下降 0.077 dB，SSIM 略有提高，说明该方法能够显著减少背景污染，并基本保持重建质量。

![初始化消融]({FIGURES["init_chart"]})

### 5.2 视角数量与覆盖范围实验

| 方法 | PSNR↑ | SSIM↑ | LPIPS↓ | 前景PSNR↑ | 前景裁剪LPIPS↓ |
|---|---:|---:|---:|---:|---:|
| 8 视角，包含少量顶部视角 | 17.958 | 0.92673 | 0.10038 | 7.522 | 0.49354 |
| 40 视角，主要为水平环绕视角 | 17.414 | 0.93085 | 0.09961 | 6.620 | 0.50889 |

结果表明，训练图像数量并不是唯一决定因素，视角的空间覆盖范围同样关键。如果训练视角主要集中在水平环绕，而缺少顶部视角，那么模型在顶部测试视角上仍会出现泛化不足。

![视角覆盖实验]({FIGURES["view_chart"]})

### 5.3 Gaussian Repair 工程验证

| 方法 | PSNR↑ | SSIM↑ | LPIPS↓ |
|---|---:|---:|---:|
| Coarse 3DGS | 28.295 | 0.97198 | 0.03663 |
| Gaussian Repair | 28.636 | 0.97145 | 0.03703 |

该实验使用已有全视角 HQ 模型作为初始化，仅用于验证 Repair 链路是否能在本地真实数据上完整运行，不属于严格少视角留出集实验。结果显示 Gaussian Repair 使拟合视角 PSNR 提高 0.342 dB，但 SSIM 和 LPIPS 略有退化，定性变化也较小。

![Repair 验证]({FIGURES["repair_chart"]})

## 六、项目难点与解决办法

| 遇到的问题 | 解决办法 | 效果 |
|---|---|---|
| 视频帧模糊、重复度高 | 使用拉普拉斯方差和时间窗口筛选关键帧 | 减少低质量输入，降低训练负担 |
| 背景点污染 COLMAP 点云 | 提出多视图掩膜一致性点云过滤 | 背景点减少约 95%，重建质量基本保持 |
| 掩膜边缘误差影响几何 | 使用颜色先验、SAM/GrabCut 与形态学后处理 | 获得更稳定的前景掩膜 |
| 真实物体纹理较少，位姿估计不稳定 | 比较 COLMAP、DUSt3R、MASt3R 等路线 | 确定更适合当前数据的位姿流程 |
| 论文流程偏向 Linux 和标准数据集 | 完成 Windows + RTX 3080 本地部署和脚本适配 | 跑通真实手机视频到新视角渲染的完整链路 |

## 七、收获与展望

通过本项目，我不仅理解了 3D Gaussian Splatting 和 GaussianObject 的基本流程，也更深刻地认识到真实数据和标准数据集之间的差距。真实手机视频中的模糊、重复、反光、低纹理、背景干扰和掩膜误差都会被三维重建算法放大，因此工程实现并不是简单地运行论文代码，而需要围绕数据质量、位姿、掩膜和初始化做大量适配。

后续工作可以进一步设计联合清晰度、相似度和空间覆盖的关键帧选择算法，并引入位姿置信度加权训练，从而提升真实场景下的重建稳定性。

## 八、结论

本课程项目完成了基于 Gaussian Splatting 的真实物体三维重建设计与实现，构建了从手机视频采集、关键帧筛选、前景分割、位姿估计、点云过滤、3DGS 训练到新视角渲染的完整流程。项目在参考 GaussianObject 的基础上，重点解决真实手机视频输入下的工程适配问题，并提出多视图掩膜约束点云过滤方法。实验表明，该方法能够将 COLMAP 初始点云减少约 95%，同时基本保持渲染指标稳定，体现出较好的实用价值。

## 参考文献

1. Yang, Chen, et al. GaussianObject: High-Quality 3D Object Reconstruction from Four Views with Gaussian Splatting. ACM Transactions on Graphics, 2024.
2. Kerbl, Bernhard, et al. 3D Gaussian Splatting for Real-Time Radiance Field Rendering. ACM Transactions on Graphics, 2023.
3. Schonberger, Johannes L., and Jan-Michael Frahm. Structure-from-Motion Revisited. CVPR, 2016.
"""
    REPORT_MD.write_text(text, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_markdown()
    print(REPORT_DOCX)
    print(REPORT_MD)
