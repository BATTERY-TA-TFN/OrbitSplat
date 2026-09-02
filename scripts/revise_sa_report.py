from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "doc_revision"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_DOCX = OUT_DIR / "SA25218120-王婧-课程报告修改版.docx"
OUT_MD = OUT_DIR / "SA25218120-王婧-课程报告修改版.md"

FIG = {
    "keyframe": ROOT / "output" / "keyframe_selection_figure.png",
    "filtering": ROOT / "output" / "colmap_mask_filtering_figure.png",
    "init_chart": ROOT / "output" / "evaluation" / "initialization_ablation_metrics.png",
    "init_qual": ROOT / "output" / "evaluation" / "initialization_ablation_qualitative.png",
    "view_chart": ROOT / "output" / "evaluation" / "common_subset_metrics.png",
    "view_qual": ROOT / "output" / "evaluation" / "common_subset_qualitative.png",
    "repair_chart": ROOT / "output" / "evaluation" / "repair_engineering_validation.png",
}


def set_run(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_title(document, text, size=18, bold=True):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run(r, "黑体", size, bold)
    return p


def add_heading(document, text, level=1):
    p = document.add_heading(text, level=level)
    for r in p.runs:
        set_run(r, "黑体", 14 if level == 1 else 12, True)
    return p


def add_para(document, text, first_line=True):
    p = document.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(3)
    if first_line:
        p.paragraph_format.first_line_indent = Pt(21)
    r = p.add_run(text)
    set_run(r)
    return p


def set_cell(cell, text, bold=False, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text))
    set_run(r, size=9, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, True, "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            set_cell(cells[i], v)
    document.add_paragraph()
    return table


def add_captioned_image(document, path, caption, width=15.0):
    if not path.exists():
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run(r, size=9, color=(80, 80, 80))


def add_code(document, text):
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.right_indent = Pt(18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, "Consolas", 8.5)


def add_static_toc(document):
    add_heading(document, "目录", 1)
    items = [
        "摘要",
        "1 引言",
        "2 实验环境与数据",
        "3 方法设计",
        "3.1 总体技术路线",
        "3.2 关键帧选择与真实视频预处理",
        "3.3 多阶段前景掩膜生成",
        "3.4 多视图掩膜约束点云过滤",
        "4 实验设计与结果分析",
        "4.1 评价指标",
        "4.2 初始化方式消融实验",
        "4.3 视角数量与覆盖范围实验",
        "4.4 Gaussian Repair 工程验证",
        "5 工程难点与解决方案",
        "6 结论与展望",
        "参考文献",
    ]
    for item in items:
        p = document.add_paragraph()
        r = p.add_run(item)
        set_run(r, size=10.5)
    document.add_page_break()


def build_docx():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    add_title(doc, "中国科学技术大学", 16, True)
    add_title(doc, "数字图像处理课程报告", 18, True)
    doc.add_paragraph()
    add_title(doc, "基于 Gaussian Splatting 的真实物体三维重建设计与实现", 16, True)
    doc.add_paragraph()
    info = [
        "姓名：王婧",
        "学号：SA25218120",
        "实验平台：Windows + RTX 3080",
        "报告日期：2026 年 6 月",
    ]
    for line in info:
        add_title(doc, line, 11, False)
    doc.add_page_break()

    add_static_toc(doc)

    add_heading(doc, "摘要", 1)
    add_para(
        doc,
        "针对真实手机视频条件下物体三维重建面临的图像质量不稳定、背景特征干扰、相机位姿估计误差和少视角泛化不足等问题，本文基于 GaussianObject 与 3D Gaussian Splatting 框架，构建了一套面向普通手机环绕视频的真实物体重建流程。主要工作包括真实视频关键帧筛选、多阶段前景掩膜生成、COLMAP/DUSt3R/MASt3R 位姿估计路线比较、多视图掩膜约束稀疏点云过滤，以及 Gaussian Repair 链路的本地工程验证。实验结果表明，多视图掩膜约束点云过滤能够将 COLMAP 初始稀疏点由 5,727 个减少至 285 个，背景点减少约 95.02%，同时在 46 个留出测试视角上基本保持渲染质量，PSNR 仅下降 0.077 dB。进一步的视角数量实验表明，训练视角的空间覆盖对未见视角质量具有重要影响，单纯增加训练图像数量并不必然带来更好的泛化效果。"
    )
    add_para(doc, "关键词：三维重建；3D Gaussian Splatting；GaussianObject；新视角合成；点云过滤；少视角重建", False)

    add_heading(doc, "1 引言", 1)
    add_para(doc, "随着数字孪生、增强现实、虚拟展示和三维内容生成的发展，如何从低成本图像输入中构建可渲染的三维物体表示，已经成为数字图像处理与计算机视觉中的重要问题。传统三维建模通常依赖结构光扫描、摄影测量设备或人工建模流程，硬件门槛和制作成本较高。近年来，NeRF 与 3D Gaussian Splatting 等基于图像的新视角合成方法显著降低了三维重建的输入要求，使普通图像序列也能够用于物体外观与几何结构恢复。")
    add_para(doc, "然而，现有论文方法通常在标准数据集、受控采集条件或 Linux/CUDA 研究环境下进行验证，与普通用户在真实环境中使用手机拍摄物体视频仍存在明显差距。真实手机视频往往包含运动模糊、重复视角、背景纹理干扰、低纹理物体匹配困难以及自动分割误差等问题。上述因素会直接影响位姿估计、点云初始化和三维高斯优化过程，导致漂浮点、几何缺失和未见视角渲染失真。")
    add_para(doc, "因此，本文的研究目标不是单纯复现 GaussianObject 论文结果，而是围绕真实手机视频输入，研究其在 Windows + RTX 3080 环境下的工程适配方法，并通过定量与定性实验分析关键帧选择、掩膜质量、点云初始化和视角覆盖对重建质量的影响。")

    add_heading(doc, "2 实验环境与数据", 1)
    add_table(doc, ["类别", "配置"], [
        ["操作系统", "Windows"],
        ["GPU", "NVIDIA RTX 3080"],
        ["Python", "3.11"],
        ["CUDA", "11.8"],
        ["PyTorch", "2.1.2"],
        ["主要依赖", "COLMAP、OpenCV、GaussianObject、3D Gaussian Splatting、LPIPS"],
        ["实验数据", "手机环绕拍摄皮卡丘视频，共 54 个已标定位姿视角"],
    ])
    add_para(doc, "实验对象为真实桌面环境下拍摄的皮卡丘玩偶。视频经关键帧抽取、前景分割和相机位姿估计后，得到包含图像、掩膜和相机参数的数据集。严格少视角实验中选取 8 个训练视角，其余 46 个视角作为测试集；训练分辨率统一缩放为原图的 1/4，训练迭代数为 10,000。")

    add_heading(doc, "3 方法设计", 1)
    add_heading(doc, "3.1 总体技术路线", 2)
    add_para(doc, "本文以 GaussianObject 为参考框架。GaussianObject 通过结构先验和生成式修复缓解少视图物体重建中的几何缺失问题。本文保留其 3DGS 表示与 Repair 工程链路，重点针对真实手机视频输入设计前处理与点云初始化改进。总体流程包括：手机视频采集、关键帧筛选、多阶段掩膜生成、位姿估计、稀疏点云过滤、粗糙 3DGS 训练、Repair 链路验证和新视角渲染。")
    add_para(doc, "为避免贡献边界混淆，本文将 Visual Hull、Leave-One-Out 和 Gaussian Repair 视为参考论文中的基础模块；本文的主要方法贡献集中在真实数据处理链路和多视图掩膜约束点云过滤。")

    add_heading(doc, "3.2 关键帧选择与真实视频预处理", 2)
    add_para(doc, "真实手机视频中相邻帧高度相似，且部分帧存在运动模糊。若直接使用全部视频帧，冗余图像会增加训练成本，模糊图像还可能引入错误监督。本文使用拉普拉斯方差作为清晰度度量，并结合时间窗口约束与视角覆盖原则选择关键帧。该策略使输入视图既具有较高图像质量，又覆盖侧面、背面和顶部等互补观察方向。")
    add_captioned_image(doc, FIG["keyframe"], "图 1 关键帧选择与清晰度筛选示意", 14.6)

    add_heading(doc, "3.3 多阶段前景掩膜生成", 2)
    add_para(doc, "前景掩膜在本文流程中不仅用于图像分割，还参与 Visual Hull 初始化、轮廓约束和点云过滤。为提升真实环境下掩膜的稳定性，本文采用由粗到精的多阶段分割策略：首先利用颜色先验定位目标区域，随后结合 GrabCut 或 SAM 获得精细掩膜，最后通过最大连通域保留、孔洞填充和形态学操作清理边缘噪声。")

    add_heading(doc, "3.4 多视图掩膜约束点云过滤", 2)
    add_para(doc, "在真实桌面拍摄场景中，COLMAP 的稀疏点云通常包含大量桌面与背景点。由于 3DGS 会以稀疏点云作为几何初始化，这些背景点会成为后续优化中的污染源。为此，本文提出多视图掩膜约束点云过滤方法。对于 COLMAP 得到的三维点 X，将其投影到多个相机视图中，并统计该点投影位置落入前景掩膜的支持次数。仅当该点在多个可见视图中满足前景一致性时，才被保留为物体初始化点。")
    add_para(doc, "形式化地，设第 i 个视图的前景掩膜为 M_i，相机投影函数为 π_i(X)，则三维点 X 的前景支持度可表示为 s(X)=Σ_i 1[M_i(π_i(X))=1]。当 s(X)/N_visible 大于给定阈值时保留该点，否则认为其更可能属于背景或错误匹配点。")
    add_code(doc, "Input: COLMAP sparse points P, camera parameters C, foreground masks M\nOutput: filtered object point cloud P_obj\n\nfor X in P:\n    support, visible = 0, 0\n    for camera i in C:\n        u = project(X, camera_i)\n        if u is inside image boundary:\n            visible += 1\n            if M_i[u] == foreground:\n                support += 1\n    if visible > 0 and support / visible >= threshold:\n        keep X in P_obj")
    add_captioned_image(doc, FIG["filtering"], "图 2 原始 COLMAP 点云、掩膜过滤点云与重建结果对比", 15.0)

    add_heading(doc, "4 实验设计与结果分析", 1)
    add_heading(doc, "4.1 评价指标", 2)
    add_para(doc, "本文采用 PSNR、SSIM 和 LPIPS 作为主要评价指标。其中 PSNR 反映像素误差，SSIM 衡量结构相似性，LPIPS 衡量感知距离。考虑到实验图像背景区域较大，本文进一步计算前景区域 PSNR 和前景裁剪 LPIPS，以更直接评估物体主体区域的重建质量。")

    add_heading(doc, "4.2 初始化方式消融实验", 2)
    add_table(doc, ["初始化方法", "初始点数", "PSNR↑", "SSIM↑", "LPIPS↓", "前景PSNR↑", "前景裁剪LPIPS↓"], [
        ["原始 COLMAP", "5,727", "26.596", "0.96273", "0.04459", "16.314", "0.26980"],
        ["掩膜过滤 COLMAP", "285", "26.519", "0.96276", "0.04515", "16.246", "0.26907"],
        ["Visual Hull", "7,787", "24.713", "0.95745", "0.05733", "13.288", "0.31887"],
    ])
    add_para(doc, "从表中可以看出，多视图掩膜过滤将初始化点数从 5,727 降至 285，仅保留约 4.98% 的点，说明背景污染被显著抑制。与此同时，PSNR 仅从 26.596 降至 26.519，下降幅度为 0.077 dB，SSIM 还略有提升。这表明过滤后的点云虽然规模大幅降低，但仍保留了对物体重建有效的几何支撑。前景裁剪 LPIPS 从 0.26980 降至 0.26907，也说明主体区域感知质量没有退化。")
    add_para(doc, "Visual Hull 初始化在当前真实数据上表现较差，PSNR 为 24.713，LPIPS 为 0.05733。其原因可能是 Visual Hull 对掩膜边界和相机位姿误差较敏感，当真实视频中的掩膜和位姿存在误差时，体素求交过程会放大这些误差，导致初始化几何不稳定。")
    add_captioned_image(doc, FIG["init_chart"], "图 3 初始化方式消融实验指标对比", 15.0)
    add_captioned_image(doc, FIG["init_qual"], "图 4 不同初始化方式在未见视角上的定性结果", 15.0)

    add_heading(doc, "4.3 视角数量与覆盖范围实验", 2)
    add_table(doc, ["方法", "PSNR↑", "SSIM↑", "LPIPS↓", "前景PSNR↑", "前景裁剪LPIPS↓"], [
        ["8 视角，包含少量顶部视角", "17.958", "0.92673", "0.10038", "7.522", "0.49354"],
        ["40 视角，主要为水平环绕视角", "17.414", "0.93085", "0.09961", "6.620", "0.50889"],
    ])
    add_para(doc, "为公平比较不同训练视角设置，本文选取 8 视角模型和 40 视角模型均未见过的 12 个顶部视角作为共同测试集。实验结果显示，40 视角模型在 SSIM 和 LPIPS 上略优，说明其整体结构一致性和感知相似性略有改善；但 8 视角模型在 PSNR 和前景指标上更高，原因是其训练集中包含少量顶部视角，而 40 视角设置主要集中于水平环绕视角。")
    add_para(doc, "该实验说明，少视角重建并不只依赖训练图像数量，还强烈依赖训练视角在三维空间中的覆盖分布。对于手机视频重建任务，关键帧选择应同时考虑清晰度、去重和空间覆盖，而不是简单选取更多连续帧。")
    add_captioned_image(doc, FIG["view_chart"], "图 5 共同未见顶部视角上的指标对比", 15.0)
    add_captioned_image(doc, FIG["view_qual"], "图 6 共同未见顶部视角上的定性对比", 15.0)

    add_heading(doc, "4.4 Gaussian Repair 工程验证", 2)
    add_table(doc, ["方法", "PSNR↑", "SSIM↑", "LPIPS↓"], [
        ["Coarse 3DGS", "28.295", "0.97198", "0.03663"],
        ["Gaussian Repair", "28.636", "0.97145", "0.03703"],
    ])
    add_para(doc, "本文还对 Gaussian Repair 链路进行了本地工程验证。需要说明的是，该实验使用已有全视角 HQ 模型作为初始化，因此不属于严格少视角留出集实验，不能直接作为未见视角泛化能力的证据。结果显示，Repair 后拟合视角 PSNR 提高 0.342 dB，但 SSIM 和 LPIPS 略有退化，定性变化较小。该结果表明 Repair 链路能够在本地真实数据上运行，但其稳定收益仍需在严格 8 视角设置下重新训练 Leave-One-Out 和 LoRA Repair 模型后进一步验证。")
    add_captioned_image(doc, FIG["repair_chart"], "图 7 Gaussian Repair 工程验证指标对比", 15.0)

    add_heading(doc, "5 工程难点与解决方案", 1)
    add_table(doc, ["问题", "技术处理", "作用"], [
        ["视频帧模糊与视角冗余", "基于拉普拉斯方差和时间窗口进行关键帧筛选", "提高输入质量，减少重复视角"],
        ["背景点污染初始化点云", "多视图掩膜一致性过滤 COLMAP 稀疏点", "显著降低背景点对 3DGS 初始化的影响"],
        ["自动分割边界不稳定", "颜色先验、SAM/GrabCut 与形态学后处理结合", "提高掩膜完整性和稳定性"],
        ["低纹理物体位姿估计困难", "对比 COLMAP、DUSt3R、MASt3R 多条路线", "选择更适合当前数据的位姿结果"],
        ["论文环境与本地环境差异", "完成 Windows + RTX 3080 脚本适配与依赖配置", "跑通从手机视频到新视角渲染的完整流程"],
    ])

    add_heading(doc, "6 结论与展望", 1)
    add_para(doc, "本文完成了基于 Gaussian Splatting 的真实物体三维重建设计与实现，构建了从手机视频采集、关键帧筛选、前景分割、位姿估计、点云过滤、3DGS 训练到新视角渲染的完整流程。实验表明，多视图掩膜约束点云过滤能够有效减少 COLMAP 背景点对物体重建的污染，在初始点数量降低约 95.02% 的情况下仍基本保持渲染质量。")
    add_para(doc, "本文还通过视角数量与覆盖范围实验说明，真实手机视频重建中关键帧的空间分布比简单增加图像数量更加重要。Gaussian Repair 工程验证表明该链路可以在本地真实数据上运行，但严格少视角 Repair 效果仍需后续进一步验证。未来工作可以从三个方向展开：第一，设计联合清晰度、相似度和空间覆盖的关键帧选择策略；第二，引入位姿置信度加权训练，降低错误位姿对优化的影响；第三，在严格少视角设置下重新训练 Repair 模型，验证生成式修复对未见视角泛化的实际贡献。")

    add_heading(doc, "参考文献", 1)
    refs = [
        "Yang, Chen, et al. GaussianObject: High-Quality 3D Object Reconstruction from Four Views with Gaussian Splatting. ACM Transactions on Graphics, 2024.",
        "Kerbl, Bernhard, et al. 3D Gaussian Splatting for Real-Time Radiance Field Rendering. ACM Transactions on Graphics, 2023.",
        "Schonberger, Johannes L., and Jan-Michael Frahm. Structure-from-Motion Revisited. CVPR, 2016.",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(ref)
        set_run(r)

    doc.save(OUT_DOCX)


def build_md():
    OUT_MD.write_text(
        "# SA25218120-王婧-课程报告修改版\n\n"
        "已根据原 .doc 报告内容重建为更规范的学术版 Word 文档。主要修正：\n\n"
        "- 删除原文件中残留的旧模板目录内容。\n"
        "- 统一章节为摘要、引言、实验环境、方法、实验分析、工程难点、结论与参考文献。\n"
        "- 补强多视图掩膜约束点云过滤的方法定义、伪代码与实验分析。\n"
        "- 保留并整理最终实验表格与关键展示图片。\n",
        encoding="utf-8",
    )


if __name__ == "__main__":
    build_docx()
    build_md()
    print(OUT_DOCX)
    print(OUT_MD)
